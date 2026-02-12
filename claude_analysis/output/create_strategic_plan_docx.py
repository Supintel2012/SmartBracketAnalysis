#!/usr/bin/env python3
"""
Generate SmartBracket 2026 Strategic Plan Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_strategic_plan():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ============ TITLE PAGE ============
    # Add some spacing at top
    for _ in range(3):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SmartBracket™')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0xF9, 0x55, 0x39)  # Orange color

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('March Madness 2026')
    run.font.size = Pt(28)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('Strategic Execution Plan')
    run.font.size = Pt(24)
    run.italic = True

    # Add spacing
    for _ in range(4):
        doc.add_paragraph()

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('Dan Lipsy')
    run.font.size = Pt(14)
    run.bold = True

    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = role.add_run('President & COO')
    run.font.size = Pt(12)

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run('Supported Intelligence LLC')
    run.font.size = Pt(12)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run('January 2026')
    run.font.size = Pt(12)

    # Add spacing and document number
    for _ in range(4):
        doc.add_paragraph()

    doc_num = doc.add_paragraph()
    doc_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_num.add_run('SI Strategic Document 2026-2')
    run.font.size = Pt(10)
    run.italic = True

    # Page break
    doc.add_page_break()

    # ============ EXECUTIVE SUMMARY ============
    doc.add_heading('Executive Summary', level=1)

    doc.add_paragraph(
        'Based on comprehensive analysis of SmartBracket\'s codebase, 9-year performance history, '
        'and market dynamics, this plan recommends a fundamental repositioning for 2026:'
    )

    p = doc.add_paragraph()
    run = p.add_run('Stop treating SmartBracket as a revenue business. Start treating it as your most compelling sales asset.')
    run.bold = True
    run.italic = True

    doc.add_paragraph(
        'SmartBracket\'s 9-year track record of ~90th percentile performance is worth more as a '
        'credibility engine for Supported Intelligence\'s enterprise RRToolbox offerings than as a '
        '$2,107/5-year consumer product.'
    )

    doc.add_heading('Strategic Priorities', level=2)

    # Priority table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Priority', 'Objective', 'Investment']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    priorities = [
        ('1', 'MDP vs. LLM Bake-Off 2.0 (Credibility Campaign)', '$2,000-5,000'),
        ('2', 'Supported Intelligence Website Integration', '$0 (student labor)'),
        ('3', 'Controlled Marketing Experiment', '$1,000-2,000'),
        ('4', 'Enterprise Lead Generation', '$0-1,000'),
    ]

    for i, (priority, objective, investment) in enumerate(priorities, 1):
        table.rows[i].cells[0].text = priority
        table.rows[i].cells[1].text = objective
        table.rows[i].cells[2].text = investment

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Total Recommended 2026 Budget: $3,000-8,000')
    run.bold = True
    p.add_run(' (vs. $9,461 spent on failed paid media 2022-2024)')

    doc.add_page_break()

    # ============ SECTION 1: STRATEGIC REPOSITIONING ============
    doc.add_heading('1. Strategic Repositioning', level=1)

    doc.add_heading('1.1 The Problem with Current Approach', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Metric', 'Historical (2020-2024)', 'Implication']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    data = [
        ('Total Revenue', '$2,107', 'Product not viable at scale'),
        ('Total Marketing Spend', '$16,897', '8x revenue spent on acquisition'),
        ('Customer Acquisition Cost', '$37.98-$94.08', 'Losing $33-89 per customer'),
        ('Consumer Market Trend', 'Commoditizing (free AI)', 'Getting worse, not better'),
    ]

    for i, (metric, value, implication) in enumerate(data, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = value
        table.rows[i].cells[2].text = implication

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Continuing the current approach will lose more money.')
    run.bold = True

    doc.add_heading('1.2 The New Positioning', level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Old Positioning'
    table.rows[0].cells[1].text = 'New Positioning'
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    changes = [
        ('SmartBracket is a consumer product', 'SmartBracket is a demonstration'),
        ('Goal: Acquire paying users', 'Goal: Generate enterprise leads'),
        ('Success metric: Revenue', 'Success metric: Qualified leads for RRToolbox'),
        ('Marketing: Paid acquisition', 'Marketing: Content + credibility'),
    ]

    for i, (old, new) in enumerate(changes, 1):
        table.rows[i].cells[0].text = old
        table.rows[i].cells[1].text = new

    doc.add_paragraph()

    doc.add_heading('1.3 Why This Works', level=2)

    reasons = [
        'The 9-year track record cannot be bought – It\'s your most defensible asset',
        'Free AI makes consumer market unprofitable – But increases interest in "beating AI"',
        'Enterprise customers want proof – SmartBracket IS the proof',
        'March Madness gets attention – Use that attention for Supported Intelligence',
    ]

    for reason in reasons:
        doc.add_paragraph(reason, style='List Bullet')

    doc.add_page_break()

    # ============ SECTION 2: BAKE-OFF CAMPAIGN ============
    doc.add_heading('2. MDP vs. LLM Bake-Off Campaign 2026', level=1)

    doc.add_heading('2.1 Campaign Concept', level=2)

    p = doc.add_paragraph()
    run = p.add_run('"MDP vs. LLM Showdown 2026"')
    run.bold = True
    run.italic = True

    doc.add_paragraph(
        'Build on the 2025 bake-off with an expanded, more visible competition that positions '
        'Supported Intelligence as the authority on optimization vs. AI.'
    )

    doc.add_heading('2.2 Bracket Categories', level=2)

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'

    headers = ['Category', 'Method', 'Purpose']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    brackets = [
        ('SmartBracket MDP', 'Pure RRToolbox optimization', 'Showcase core technology'),
        ('SmartBracket + User Input', 'MDP with 3-question adjustments', 'Show personalization value'),
        ('Claude/GPT/Gemini', 'Top 3 LLMs, identical prompt', 'Benchmark against AI hype'),
        ('Hybrid: LLM → MDP', 'LLM picks fed into MDP as locked picks', 'Novel approach, good story'),
        ('Hybrid: MDP → LLM', 'MDP recommendations given to LLM', 'Alternative hybrid'),
        ('ESPN Smart Bracket', 'ESPN\'s tool', 'Free baseline comparison'),
        ('Chalk Bracket', 'All favorites', 'Naive baseline'),
        ('Random Student Picks', 'GVSU/Turkey student brackets', 'Human baseline'),
    ]

    for i, (category, method, purpose) in enumerate(brackets, 1):
        table.rows[i].cells[0].text = category
        table.rows[i].cells[1].text = method
        table.rows[i].cells[2].text = purpose

    doc.add_paragraph()

    doc.add_heading('2.3 Campaign Timeline', level=2)

    doc.add_heading('Pre-Tournament (February 15 - Selection Sunday)', level=3)

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    headers = ['Week', 'Activity', 'Owner', 'Cost']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    pre_tournament = [
        ('Week 1', 'Publish "2025 Bake-Off Results" blog post', 'Student Team', '$0'),
        ('Week 1', 'Create 2026 methodology page', 'Student Team', '$0'),
        ('Week 2', 'Set up tracking infrastructure', 'Student Team', '$0'),
        ('Week 2', 'Recruit LLMs for competition', 'Student Team', '$0'),
        ('Week 3', 'Generate all brackets', 'Student Team', '$0'),
        ('Selection Sunday', 'Publish all brackets publicly', 'Student Team', '$0'),
    ]

    for i, (week, activity, owner, cost) in enumerate(pre_tournament, 1):
        table.rows[i].cells[0].text = week
        table.rows[i].cells[1].text = activity
        table.rows[i].cells[2].text = owner
        table.rows[i].cells[3].text = cost

    doc.add_paragraph()

    doc.add_heading('Tournament Phase (March)', level=3)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['Activity', 'Frequency', 'Owner', 'Cost']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    tournament = [
        ('Daily leaderboard updates', 'Daily', 'Student Team', '$0'),
        ('Weekly analysis posts', 'Weekly', 'Student Team', '$0'),
        ('Social media updates', 'Daily during games', 'Student Team', '$0'),
        ('Final results & analysis', 'Post-tournament', 'COO + Student Team', '$0'),
    ]

    for i, (activity, freq, owner, cost) in enumerate(tournament, 1):
        table.rows[i].cells[0].text = activity
        table.rows[i].cells[1].text = freq
        table.rows[i].cells[2].text = owner
        table.rows[i].cells[3].text = cost

    doc.add_page_break()

    # ============ SECTION 3: STUDENT RESOURCES ============
    doc.add_heading('3. Student Resource Plan', level=1)

    doc.add_heading('3.1 Available Resources', level=2)

    doc.add_paragraph(
        'Two student resource pools are available for 2026 execution:'
    )

    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'

    headers = ['Resource', 'Location', 'Strengths', 'Considerations']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    resources = [
        ('Grand Valley State University', 'Michigan, USA', 'Local, same timezone, campus resources', 'Higher hourly rate ($15-20/hr)'),
        ('International Students', 'Turkey', 'Lower cost, strong technical skills', 'Timezone difference (8 hrs), remote only'),
    ]

    for i, (resource, location, strengths, considerations) in enumerate(resources, 1):
        table.rows[i].cells[0].text = resource
        table.rows[i].cells[1].text = location
        table.rows[i].cells[2].text = strengths
        table.rows[i].cells[3].text = considerations

    doc.add_paragraph()

    doc.add_heading('3.2 Recommended Team Structure', level=2)

    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'

    headers = ['Role', 'Source', 'Hours/Week', 'Duration', 'Est. Cost']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    team = [
        ('Content Lead', 'GVSU', '8-10', 'Feb-Apr (12 wks)', '$1,440-2,400'),
        ('Social Media Coordinator', 'Turkey', '5-8', 'Feb-Apr (12 wks)', '$300-600'),
        ('Data Analyst', 'Turkey or GVSU', '5-8', 'Mar-Apr (6 wks)', '$300-960'),
    ]

    for i, (role, source, hours, duration, cost) in enumerate(team, 1):
        table.rows[i].cells[0].text = role
        table.rows[i].cells[1].text = source
        table.rows[i].cells[2].text = hours
        table.rows[i].cells[3].text = duration
        table.rows[i].cells[4].text = cost

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Total Student Labor Budget: $2,040-3,960')
    run.bold = True

    doc.add_heading('3.3 Task Allocation', level=2)

    doc.add_heading('GVSU Student (Content Lead)', level=3)

    gvsu_tasks = [
        'Write all blog posts and long-form content',
        'Create methodology documentation',
        'Coordinate with COO on messaging',
        'Quality control on all public materials',
        'Lead LinkedIn content strategy',
        'Handle any in-person campus promotion',
    ]

    for task in gvsu_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('Turkey Student(s) (Execution Support)', level=3)

    turkey_tasks = [
        'Daily social media posting (Twitter/X during games)',
        'Create graphics and infographics',
        'Update tracking spreadsheets',
        'Monitor and respond to social engagement',
        'Research competitor brackets',
        'Compile data for analysis',
    ]

    for task in turkey_tasks:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_page_break()

    # ============ SECTION 4: BUDGET ============
    doc.add_heading('4. Budget Summary', level=1)

    doc.add_heading('4.1 2026 Investment Allocation', level=2)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    headers = ['Category', 'Low Estimate', 'High Estimate']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    budget = [
        ('GVSU Student Labor', '$1,440', '$2,400'),
        ('Turkey Student Labor', '$600', '$1,560'),
        ('LinkedIn Promoted Posts', '$300', '$500'),
        ('Content/Design Tools', '$200', '$500'),
        ('Miscellaneous', '$200', '$400'),
        ('TOTAL', '$2,740', '$5,360'),
    ]

    for i, (category, low, high) in enumerate(budget, 1):
        table.rows[i].cells[0].text = category
        table.rows[i].cells[1].text = low
        table.rows[i].cells[2].text = high
        if category == 'TOTAL':
            for j in range(3):
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('4.2 Comparison to Historical Spend', level=2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ['Year', 'Spend', 'Result']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    history = [
        ('2022', '$5,750', '49 users, $117 CAC'),
        ('2023-24', '$9,461+', 'Negative ROI'),
        ('2026 (recommended)', '$2,740-5,360', 'Enterprise leads + credibility'),
    ]

    for i, (year, spend, result) in enumerate(history, 1):
        table.rows[i].cells[0].text = year
        table.rows[i].cells[1].text = spend
        table.rows[i].cells[2].text = result

    doc.add_paragraph()

    doc.add_heading('4.3 ROI Scenario Analysis', level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['Scenario', 'Probability', 'Enterprise Value', 'ROI']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    scenarios = [
        ('1 enterprise deal closes', '30%', '$50,000-500,000', '10x-100x'),
        ('2-3 qualified leads', '50%', '$10,000-50,000 pipeline', '2x-10x'),
        ('Credibility only', '90%', 'Hard to quantify', 'Positive'),
        ('Complete failure', '10%', '$0', '-100%'),
    ]

    for i, (scenario, prob, value, roi) in enumerate(scenarios, 1):
        table.rows[i].cells[0].text = scenario
        table.rows[i].cells[1].text = prob
        table.rows[i].cells[2].text = value
        table.rows[i].cells[3].text = roi

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Expected Value: ')
    run.bold = True
    p.add_run('Even with 30% probability of closing one enterprise deal ($50K-500K), the expected value far exceeds the investment.')

    doc.add_page_break()

    # ============ SECTION 5: EXECUTION TIMELINE ============
    doc.add_heading('5. Execution Timeline', level=1)

    doc.add_heading('5.1 January 2026 (Preparation)', level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['Week', 'Milestone', 'Deliverable', 'Owner']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    jan = [
        ('Jan 20-26', 'Security fixes', 'MD5 replaced, secrets removed', 'Dev'),
        ('Jan 20-26', 'Student recruitment', 'Offers extended', 'COO'),
        ('Jan 27-Feb 2', 'Student onboarding', 'Access granted, training', 'COO'),
        ('Jan 27-Feb 2', 'Tool setup', 'Social accounts, analytics', 'Student Team'),
    ]

    for i, (week, milestone, deliverable, owner) in enumerate(jan, 1):
        table.rows[i].cells[0].text = week
        table.rows[i].cells[1].text = milestone
        table.rows[i].cells[2].text = deliverable
        table.rows[i].cells[3].text = owner

    doc.add_paragraph()

    doc.add_heading('5.2 February 2026 (Pre-Tournament)', level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['Week', 'Milestone', 'Deliverable', 'Owner']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    feb = [
        ('Feb 3-9', '2025 results published', 'Blog post live', 'Content Lead'),
        ('Feb 10-16', '2026 methodology finalized', 'Methodology page live', 'COO + Content'),
        ('Feb 17-23', 'Bracket generation', 'All brackets ready', 'Student Team'),
        ('Feb 24-Mar 2', 'Selection Sunday prep', 'All content staged', 'Student Team'),
    ]

    for i, (week, milestone, deliverable, owner) in enumerate(feb, 1):
        table.rows[i].cells[0].text = week
        table.rows[i].cells[1].text = milestone
        table.rows[i].cells[2].text = deliverable
        table.rows[i].cells[3].text = owner

    doc.add_paragraph()

    doc.add_heading('5.3 March 2026 (Tournament Execution)', level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['Week', 'Milestone', 'Deliverable', 'Owner']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    mar = [
        ('Selection Sunday', 'Brackets published', 'All brackets public', 'Student Team'),
        ('Round 1-2', 'Daily tracking begins', 'Daily leaderboard updates', 'Turkey Student'),
        ('Sweet 16', 'Mid-tournament analysis', 'Analysis blog post', 'Content Lead'),
        ('Final Four/Champ', 'Peak engagement', 'Real-time updates', 'Full Team'),
    ]

    for i, (week, milestone, deliverable, owner) in enumerate(mar, 1):
        table.rows[i].cells[0].text = week
        table.rows[i].cells[1].text = milestone
        table.rows[i].cells[2].text = deliverable
        table.rows[i].cells[3].text = owner

    doc.add_paragraph()

    doc.add_heading('5.4 April 2026 (Post-Tournament)', level=2)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    headers = ['Week', 'Milestone', 'Deliverable', 'Owner']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    apr = [
        ('Apr 7-13', 'Final results analysis', 'Comprehensive results report', 'COO + Content'),
        ('Apr 14-20', 'Case study creation', 'SI website updated', 'Content Lead'),
        ('Apr 21-30', 'Lead follow-up', 'Outreach to inquiries', 'COO'),
    ]

    for i, (week, milestone, deliverable, owner) in enumerate(apr, 1):
        table.rows[i].cells[0].text = week
        table.rows[i].cells[1].text = milestone
        table.rows[i].cells[2].text = deliverable
        table.rows[i].cells[3].text = owner

    doc.add_page_break()

    # ============ SECTION 6: SUCCESS METRICS ============
    doc.add_heading('6. Success Metrics', level=1)

    doc.add_heading('6.1 Primary KPIs (Enterprise Focus)', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Metric', 'Target', 'Measurement']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    primary = [
        ('Enterprise inquiries from bake-off', '5-10', 'Contact form submissions'),
        ('Qualified enterprise opportunities', '2-3', 'Passed qualification questions'),
        ('SI website traffic increase', '+50% during Mar-Apr', 'Google Analytics'),
        ('LinkedIn engagement', '1,000+ impressions', 'LinkedIn analytics'),
    ]

    for i, (metric, target, measurement) in enumerate(primary, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = target
        table.rows[i].cells[2].text = measurement

    doc.add_paragraph()

    doc.add_heading('6.2 Secondary KPIs (Demonstration Success)', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Metric', 'Target', 'Measurement']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    secondary = [
        ('SmartBracket vs. LLM performance', 'Top 3 in bake-off', 'Tournament results'),
        ('Bake-off content views', '500+', 'Blog analytics'),
        ('Email list growth', '+100 subscribers', 'MailChimp'),
        ('Media mentions', '1-2', 'Google Alerts'),
    ]

    for i, (metric, target, measurement) in enumerate(secondary, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = target
        table.rows[i].cells[2].text = measurement

    doc.add_page_break()

    # ============ SECTION 7: RISKS ============
    doc.add_heading('7. Risk Assessment', level=1)

    doc.add_heading('7.1 Risks & Mitigations', level=2)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Risk', 'Probability', 'Impact', 'Mitigation']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    risks = [
        ('SmartBracket underperforms', '30%', 'High', 'Frame as "one year" vs. 9-year track record'),
        ('No enterprise leads generated', '40%', 'Medium', 'Content still has SEO/credibility value'),
        ('Student unavailability', '20%', 'Medium', 'Have backup students identified'),
        ('Security breach (MD5)', '10%', 'Critical', 'Fix before tournament (P0 priority)'),
        ('LLMs significantly improve', '50%', 'Medium', 'Emphasize consistency, not single-year'),
    ]

    for i, (risk, prob, impact, mitigation) in enumerate(risks, 1):
        table.rows[i].cells[0].text = risk
        table.rows[i].cells[1].text = prob
        table.rows[i].cells[2].text = impact
        table.rows[i].cells[3].text = mitigation

    doc.add_paragraph()

    doc.add_heading('7.2 Contingency Plans', level=2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('If SmartBracket underperforms:')
    run.bold = True

    contingency1 = [
        'Emphasize 9-year track record vs. single year',
        'Highlight variance in LLM performance',
        'Focus on "consistency" narrative',
    ]

    for item in contingency1:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('If no enterprise leads:')
    run.bold = True

    contingency2 = [
        'Content still builds SEO and credibility',
        'Learnings inform 2027 approach',
        'Consider academic publication route',
    ]

    for item in contingency2:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============ SECTION 8: IMMEDIATE ACTIONS ============
    doc.add_heading('8. Immediate Actions', level=1)

    doc.add_paragraph(
        'The following actions should be taken immediately to execute this plan:'
    )

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['Priority', 'Action', 'Timeline', 'Owner']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    actions = [
        ('P0', 'Fix MD5 security vulnerability', 'This week', 'Dev'),
        ('P0', 'Remove hardcoded FB secret', 'This week', 'Dev'),
        ('P1', 'Recruit GVSU student (Content Lead)', 'Jan 20-27', 'COO'),
        ('P1', 'Identify Turkey student(s)', 'Jan 20-27', 'COO'),
        ('P2', 'Update SI website with SmartBracket integration', 'Feb 1-15', 'Student Team'),
    ]

    for i, (priority, action, timeline, owner) in enumerate(actions, 1):
        table.rows[i].cells[0].text = priority
        table.rows[i].cells[1].text = action
        table.rows[i].cells[2].text = timeline
        table.rows[i].cells[3].text = owner

    doc.add_paragraph()
    doc.add_paragraph()

    # Closing
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— End of Strategic Plan —')
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Prepared by Claude Code Business Analysis Framework')
    run.font.size = Pt(9)
    run.italic = True

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run('For Supported Intelligence LLC | January 2026')
    run.font.size = Pt(9)
    run.italic = True

    # Save document
    output_path = 'claude_analysis/output/SmartBracket_2026_Execution_Plan.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_strategic_plan()
